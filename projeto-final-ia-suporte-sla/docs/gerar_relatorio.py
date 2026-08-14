# -*- coding: utf-8 -*-
"""Gera o relatorio do projeto (docx) conforme item 7.3 do enunciado."""
import sys, pathlib, json
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1]))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config.settings import EVID_DIR, ROOT

M = json.load(open(EVID_DIR / "metrics.json", encoding="utf-8"))
DASH = json.load(open(ROOT / "data" / "curated" / "dashboard" / "dashboard_data.json", encoding="utf-8"))
AZUL = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def _n(v, casas=3):
    """Numero no padrao brasileiro (virgula decimal)."""
    return f"{v:.{casas}f}".replace(".", ",")


def h(txt, lvl=1):
    p = doc.add_heading(txt, level=lvl)
    for r in p.runs: r.font.color.rgb = AZUL
    return p

def para(txt):
    p = doc.add_paragraph(); p.add_run(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(txt): doc.add_paragraph(txt, style="List Bullet")

def img(path, w=6.3, cap=None):
    if pathlib.Path(path).exists():
        doc.add_picture(str(path), width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap:
            c = doc.add_paragraph(); rr = c.add_run(cap)
            rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x55,0x55,0x55)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(hd); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; rn = cells[i].paragraphs[0].add_run(str(v)); rn.font.size = Pt(9.5)
    return t

# ---------------- CAPA ----------------
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run("INSTITUTO FEDERAL DE EDUCAÇÃO, CIÊNCIA E TECNOLOGIA - GOIÁS")
r.bold = True; r.font.size = Pt(13); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); r = p.add_run("Pós-Graduação em Inteligência Artificial Aplicada - Módulo 2")
r.font.size = Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run("Pipeline de Dados em Nuvem para Aprendizagem de Máquina")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = AZUL; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); r = p.add_run("Previsão de violação de SLA em chamados de suporte técnico para apoio à decisão")
r.font.size = Pt(13); r.italic = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4): doc.add_paragraph()
for line in ["Integrantes do grupo:", "João Paulo A. S. Monteiro", "Rogério B. dos Anjos", "Pedro Felipe de Moraes Carrijo"]:
    pp = doc.add_paragraph(); rr = pp.add_run(line); rr.font.size = Pt(11)
    if line.endswith(":"): rr.bold = True
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
for line in ["Aprendizagem de Máquina - Prof. Dr. Sirlon Diniz",
             "Cloud Computing - Prof. Dr. Raphael Gomes",
             "Modelagem de Dados para IA - Prof. Dr. Otávio Calaça"]:
    pp = doc.add_paragraph(); rr = pp.add_run(line); rr.font.size = Pt(10.5)
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run("Goiânia - GO, junho de 2026"); r.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ---------------- SUMARIO ----------------
h("Sumário", 1)
par = doc.add_paragraph()
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
run = OxmlElement('w:r'); t = OxmlElement('w:t')
t.text = "Sumário automático: clique aqui com o botão direito > Atualizar campo > Atualizar o índice inteiro."
run.append(t); fld.append(run); par._p.append(fld)
doc.add_page_break()

# ---------------- 1 ----------------
h("1. Definição do problema e decisão apoiada", 1)
para("O domínio escolhido pela nossa equipe foi o de suporte técnico (service desk). Em operações "
     "de suporte, cada chamado possui um acordo de nível de serviço (SLA) de resolução, e o estouro "
     "desse prazo gera insatisfação do cliente e, muitas vezes, penalidades contratuais. O problema "
     "que atacamos é antecipar quais chamados têm maior risco de violar o SLA de resolução, de modo "
     "que o time consiga agir antes de o prazo estourar.")
para("Os elementos do problema, conforme solicitado no item 4.1, são:")
table(["Elemento", "Definição no nosso projeto"], [
    ["Domínio de aplicação", "Suporte técnico / service desk de uma empresa de software (SaaS)"],
    ["Usuário / tomador de decisão", "Gestor (coordenador) da equipe de suporte"],
    ["Decisão apoiada", "Para onde realocar agentes e quais chamados priorizar na fila"],
    ["Fontes de dados", "Tickets (estruturado), descrições textuais (não estruturado) e logs (JSON)"],
    ["Tarefa de ML", "Classificação binária: prever sla_breach (violou ou não o SLA)"],
    ["Resultado esperado", "Fila priorizada por risco e indicadores para a gestão do suporte"],
])
para("A decisão é tomada logo após a primeira resposta ao cliente. Esse recorte é importante porque "
     "define quais atributos podemos usar sem cometer vazamento de dados: usamos apenas o que já é "
     "conhecido nesse momento e nunca o tempo de resolução em si, que é justamente o desfecho que "
     "queremos antecipar.")

# ---------------- 2 ----------------
h("2. Conjuntos de dados utilizados", 1)
para("Atendendo ao requisito de combinar dados estruturados com pelo menos uma fonte não estruturada "
     "ou semiestruturada, trabalhamos com três fontes complementares. Optamos por gerar os dados de "
     "forma sintética e reprodutível (com semente fixa), por questões de privacidade e para garantir "
     "que qualquer pessoa consiga reproduzir o pipeline sem depender de credenciais externas. O esquema "
     "foi desenhado espelhando o Customer Support Ticket Dataset citado no enunciado; para usar dados "
     "reais, basta substituir os arquivos da pasta data/raw mantendo os mesmos campos.")
table(["Fonte", "Tipo", "Formato", "Registros", "Conteúdo principal"], [
    ["tickets.csv", "Estruturado", "CSV", "8.000", "categoria, prioridade, canal, tier, tempos, CSAT"],
    ["ticket_messages.jsonl", "Não estruturado", "JSONL", "8.000", "texto livre do cliente + nota do agente"],
    ["ticket_events.json", "Semiestruturado", "JSON", "8.000", "trilha de eventos/status de cada chamado"],
])
para("A variável-alvo (sla_breach) foi definida a partir do SLA de resolução por prioridade "
     "(Crítica: 4h; Alta: 8h; Média: 24h; Baixa: 48h). A prevalência de violações no conjunto ficou "
     f"em torno de {DASH['kpis']['taxa_violacao_pct']:.1f}%, o que configura um problema desbalanceado, "
     "como é comum na prática. Documentamos origem, formato, campos e limitações dos dados no arquivo "
     "docs/dicionario_de_dados.md do repositório.")

# ---------------- 3 ----------------
h("3. Arquitetura geral da solução", 1)
para("A solução foi desenvolvida com uma arquitetura em camadas (bronze/prata/ouro), seguindo a "
     "abordagem ELT (Extract, Load, Transform). No ambiente de desenvolvimento, usamos o DuckDB como "
     "proxy local do Snowflake (a mesma modelagem dbt roda nos dois, bastando trocar o profile), "
     "Airflow para orquestração, Python para processamento/ML e Metabase para o dashboard. O Amazon "
     "S3 é usado para o armazenamento bruto em nuvem.")
para("Como pedido no item 4.5, também desenhamos a arquitetura equivalente 100% em serviços AWS, "
     "apresentada na figura a seguir.")
img(ROOT / "infra" / "arquitetura_aws.png", 6.6, "Figura 1. Arquitetura equivalente 100% em AWS.")
para("Nessa arquitetura, a ingestão roda em AWS Lambda, os dados ficam no S3 em três camadas (raw, "
     "processed, curated), a carga e a modelagem analítica acontecem no Snowflake com dbt, o "
     "treino/inferência do modelo fica no Amazon SageMaker, a orquestração usa o Amazon MWAA (Airflow "
     "gerenciado), o dashboard roda em Metabase sobre EC2/ECS, e o monitoramento e a segurança ficam a "
     "cargo de CloudWatch e IAM/KMS. Toda a infraestrutura base é descrita como código no template "
     "CloudFormation (infra/cloudformation.yaml).")

# ---------------- 4 ----------------
h("4. Processamento e extração de atributos", 1)
para("Para os dados estruturados, fizemos limpeza, padronização de categorias, tratamento de valores "
     "ausentes (o CSAT ausente virou uma flag, sem imputar o alvo) e derivamos partes de data (hora, "
     "dia da semana, mês) úteis para análise.")
para("Para o texto livre (não estruturado), extraímos atributos como comprimento do texto, número de "
     "palavras, quantidade de pontos de interrogação e exclamação, proporção de letras maiúsculas, "
     "contagem de termos de um léxico de urgência e flags de palavras-chave (reembolso, cancelamento, "
     "erro). Para os logs (semiestruturado), agregamos o número de eventos, o número de mudanças de "
     "status e se houve reabertura.")
para("O resultado é uma tabela analítica única (mart_ml_features), com 25 atributos por chamado mais "
     "o alvo, pronta para o modelo. Reforçamos que o tempo de resolução e o CSAT foram deliberadamente "
     "deixados de fora dos atributos para evitar vazamento de dados.")

# ---------------- 5 ----------------
h("5. Pipeline de ELT (Airflow, dbt e Snowflake)", 1)
para("O pipeline é orquestrado por uma DAG do Airflow que encadeia as etapas na ordem: ingestão, "
     "processamento/atributos, treino e avaliação do modelo, carga no warehouse, build do dbt (modelos "
     "mais testes) e export do dashboard. Para facilitar a execução local, também fornecemos um "
     "run_pipeline.py que roda exatamente a mesma sequência sem precisar do Airflow.")
para("No dbt, organizamos os modelos em camadas:")
bullet("Staging (views): stg_tickets, stg_text_features, stg_log_features e stg_predictions;")
bullet("Dimensões: dim_channel, dim_product, dim_category e dim_date;")
bullet("Fatos e marts: fct_tickets, mart_ml_features, fct_predictions e mart_ticket_decision.")
para("Implementamos 10 testes de qualidade de dados no dbt (acima do mínimo de 2 pedido), incluindo "
     "not_null e unique em chaves, accepted_values em prioridade e no alvo, e um teste de relationships "
     "entre as predições e o fato de chamados. Na última execução, todos os testes passaram (PASS=22, "
     "0 erros). A documentação dos modelos está nos arquivos .yml do projeto dbt e pode ser publicada "
     "com dbt docs.")

# ---------------- 6 ----------------
h("6. Uso de recursos da AWS", 1)
para("Conforme o item 4.5, a solução utiliza o Amazon S3 para o armazenamento dos dados, organizados "
     "em três camadas (raw, processed, curated). A estrutura completa de nuvem foi descrita como código "
     "no template CloudFormation, que provisiona:")
table(["Recurso AWS", "Papel na solução"], [
    ["S3 (3 buckets)", "Data lake em camadas raw/processed/curated, com versionamento e criptografia"],
    ["IAM Role", "Permissões mínimas de acesso ao S3 e ao CloudWatch"],
    ["Lambda", "Função de ingestão (stub) que grava dados na camada raw"],
    ["EventBridge", "Agendamento diário da ingestão"],
    ["CloudWatch Logs", "Monitoramento e retenção de logs do pipeline"],
])
para("O pipeline inclui uma etapa de sincronização com o S3 (src/cloud/s3_sync.py, via boto3) que "
     "envia as camadas raw, processed e curated para o bucket quando as credenciais AWS e a variável "
     "S3_BUCKET estão definidas; sem elas, a etapa é pulada e a execução continua local. Validamos a "
     "lógica de upload com um S3 simulado (biblioteca moto), confirmando o envio das três camadas. "
     "Assim, o uso do S3 está implementado, bastando apontar o bucket gerado pelo CloudFormation.")
para("Consideramos também a execução no AWS Academy (Learner Lab), ambiente usado no curso, que "
     "restringe a criação de papéis IAM (permite apenas o papel pré-existente LabRole), trava a região "
     "em us-east-1 e não disponibiliza serviços gerenciados como o MWAA. Para esse cenário, "
     "disponibilizamos uma variante do template (infra/cloudformation-academy.yaml) que reutiliza o "
     "LabRole em vez de criar um papel, e utilizamos o upload para o S3, etapa que não exige papel "
     "algum (basta as credenciais temporárias do lab). Dessa forma, o requisito de usar um serviço da "
     "AWS é cumprido mesmo sob as restrições do Academy, enquanto a arquitetura completa (MWAA, "
     "SageMaker) permanece como proposta de referência no diagrama. As instruções estão em "
     "docs/aws_academy.md.")

# ---------------- 7 ----------------
h("7. Aprendizagem de Máquina: modelos e métricas", 1)
para("A tarefa é de classificação binária. Seguindo o item 4.6, criamos baselines, implementamos o "
     "modelo na forma hard-code e depois com biblioteca, e comparamos os resultados.")
para("Usamos os algoritmos vistos na disciplina: KNN e MLP. Implementamos o KNN do zero "
     "(apenas NumPy): ele guarda os exemplos e, para prever, calcula a distância euclidiana até todos os "
     "pontos de treino, seleciona os k vizinhos mais próximos (usamos k = 15) e decide pela votação da "
     "maioria. Em seguida rodamos o mesmo KNN com o scikit-learn, e as duas versões ficaram idênticas "
     "(concordância de 100% nas predições), o que nos dá confiança de que a implementação manual está "
     "correta. Como comparação, treinamos ainda um MLP — uma rede neural com duas camadas ocultas de 64 e "
     "32 neurônios. O MLP teve o melhor desempenho (F1 = " + _n(M["metricas"]["mlp"]["f1"])
     + "; recall = " + _n(M["metricas"]["mlp"]["recall"])
     + "; ROC-AUC = " + _n(M["metricas"]["mlp"]["roc_auc"])
     + ") e foi escolhido como modelo de produção.")
mm = M["metricas"]
def row(k, nome):
    d = mm[k]
    return [nome] + [_n(d[c]) for c in ("accuracy", "precision", "recall", "f1", "roc_auc")]
table(["Modelo", "Acurácia", "Precisão", "Recall", "F1", "ROC-AUC"], [
    row("baseline_majoritaria", "Baseline (classe majoritária)"),
    row("baseline_regra_prioridade", "Baseline (regra por prioridade)"),
    row("knn_hardcode", "KNN (hard-code)"),
    row("knn_sklearn", "KNN (scikit-learn)"),
    row("mlp", "MLP (produção)"),
])
para("As métricas foram escolhidas conforme a tarefa (classificação): acurácia, precisão, recall, F1 "
     "e ROC-AUC, além da matriz de confusão. Demos atenção especial ao recall, pois deixar de sinalizar "
     "um chamado que vai violar o SLA (falso negativo) é mais custoso para a operação do que um alarme "
     "falso.")
img(EVID_DIR / "matriz_confusao.png", 3.3)
img(EVID_DIR / "curva_roc.png", 3.6, "Figura 2. Matriz de confusão e curva ROC (conjunto de teste).")

# ---------------- 8 ----------------
h("8. Dashboard e principais análises", 1)
para("O dashboard foi construído no Metabase sobre as tabelas do schema analytics. Ele reúne os "
     "indicadores principais (volume de chamados, taxa de violação de SLA, tempo médio de resolução e "
     "CSAT), a visualização dos dados tratados (violação por prioridade, categoria, canal e mês), os "
     "resultados do modelo (matriz de confusão e distribuição de risco) e uma fila priorizada pelos "
     "chamados de maior risco. Há filtros por prioridade, categoria, canal e mês.")
img(EVID_DIR / "dashboard_mockup.png", 6.6, "Figura 3. Painel de apoio à decisão (visão consolidada).")
para("As análises mostram que as categorias Solicitação de Recurso e Bug/Defeito concentram as maiores "
     "taxas de violação, e que prioridade Crítica e Alta são as mais críticas. Na prática, o gestor usa "
     "a fila priorizada por risco para realocar agentes antes de o prazo estourar, e os gráficos por "
     "categoria/canal para decisões estruturais (reforço de equipe, base de conhecimento, automações).")

# ---------------- 9 ----------------
h("9. Avaliação do pipeline e análise qualitativa", 1)
para("Além das métricas do modelo, avaliamos o pipeline como um todo: as etapas executam na ordem "
     "correta, os testes do dbt garantem a qualidade dos dados após a transformação, e existe "
     "rastreabilidade do dado bruto até a predição (raw -> processed -> marts -> predictions).")
h("Análise qualitativa", 2)
bullet("Acertos: chamados de prioridade Alta/Crítica em categorias demoradas (Bug/Defeito, Solicitação "
       "de Recurso) são corretamente sinalizados com alta probabilidade de violação.")
bullet("Erros: o modelo erra mais em casos de fronteira (prioridade Média com texto curto), gerando "
       "alguns falsos negativos.")
bullet("Possíveis causas: sobreposição natural entre as classes e ausência de atributos sobre a carga "
       "do agente específico e o histórico do cliente.")
bullet("Limitações dos dados: por serem sintéticos, não capturam toda a complexidade de textos reais "
       "(gírias, erros de digitação, múltiplos idiomas).")
bullet("Riscos de uso: o modelo pode reforçar vieses operacionais (priorizar sempre os mesmos perfis) "
       "e há risco de uso do score como única base de decisão.")
bullet("Melhorias futuras: incorporar TF-IDF/embeddings do texto, adicionar features de carga do time, "
       "calibrar o threshold por custo de negócio e monitorar drift do modelo.")

# ---------------- 10 ----------------
h("10. Limitações e próximos passos", 1)
para("A camada de nuvem já está implementada e plugável: o S3 recebe as camadas de dados via boto3, "
     "a infraestrutura é descrita como código em CloudFormation, o dbt possui profile para Snowflake e a "
     "orquestração tem uma DAG de Airflow. As principais limitações são o uso de dados sintéticos e o fato "
     "de a integração de nuvem ter sido validada localmente (S3 testado com mock; Snowflake, MWAA e "
     "SageMaker dependem de contas para execução em produção). Como próximos passos de escala, pretendemos: "
     "(1) executar a carga em um Snowflake real via o profile já preparado; (2) subir a DAG no Amazon MWAA; "
     "(3) treinar e servir o modelo no SageMaker; e (4) evoluir os atributos textuais com técnicas de PLN. "
     "Mesmo assim, a solução já cumpre o ciclo completo pedido no enunciado, da ingestão à visualização "
     "para a tomada de decisão.")

# ---------------- REF ----------------
h("Referências", 1)
for ref in [
    "ENUNCIADO. Projeto Final Integrado: Pipeline de Dados em Nuvem para Aprendizagem de Máquina. IFG, Pós-Graduação em IA Aplicada, Módulo 2, 2026.",
    "dbt Labs. dbt documentation. Disponível em: https://docs.getdbt.com.",
    "Snowflake. Snowflake Documentation. Disponível em: https://docs.snowflake.com.",
    "Apache Airflow. Documentation. Disponível em: https://airflow.apache.org/docs.",
    "Amazon Web Services. AWS Documentation (S3, Lambda, CloudFormation, SageMaker, MWAA). Disponível em: https://docs.aws.amazon.com.",
    "PEDREGOSA, F. et al. Scikit-learn: Machine Learning in Python. JMLR, v. 12, 2011.",
    "Metabase. Metabase Documentation. Disponível em: https://www.metabase.com/docs.",
]:
    doc.add_paragraph(ref, style="List Bullet")

# pede ao Word para atualizar os campos (sumario) ao abrir o documento
settings = doc.settings.element
upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
settings.insert(0, upd)

out = ROOT / "docs" / "Relatorio_Projeto_Final.docx"
doc.save(str(out))
print("relatorio (acentuado) salvo:", out)
